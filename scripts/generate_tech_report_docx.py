from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "personal_tech_report_recommendation.md"
OUTPUT = ROOT / "output" / "doc" / "TR-X조-이름-학번.docx"


def clean_text(value: str) -> str:
    return value.replace("`", "").strip()


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = document.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(10.5)


def add_title(document: Document, title: str, subtitle: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.italic = True
    run2.font.size = Pt(11)


def add_markdown_line(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        document.add_paragraph("")
        return

    if stripped.startswith("## "):
        heading = clean_text(stripped[3:])
        if heading == "개인 Tech Report":
            return
        document.add_heading(heading, level=1)
        return

    if stripped.startswith("### "):
        document.add_heading(clean_text(stripped[4:]), level=2)
        return

    if stripped.startswith("- "):
        document.add_paragraph(clean_text(stripped[2:]), style="List Bullet")
        return

    numbered = stripped.split(". ", 1)
    if len(numbered) == 2 and numbered[0].isdigit():
        document.add_paragraph(clean_text(numbered[1]), style="List Number")
        return

    if stripped.startswith("|") and stripped.endswith("|"):
        # Tables are handled in a separate pass.
        return

    document.add_paragraph(clean_text(stripped))


def flush_table(document: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return

    header = rows[0]
    body = rows[2:] if len(rows) >= 3 and all(cell.startswith("---") for cell in rows[1]) else rows[1:]
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"

    for index, value in enumerate(header):
        table.rows[0].cells[index].text = clean_text(value)

    for row in body:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = clean_text(value)


def build_document() -> Document:
    document = Document()
    configure_document(document)
    add_title(document, "멀티모달 검색 및 Multi-Stage 추천 시스템 구축", "개인 Tech Report")

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    table_buffer: list[list[str]] = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            row = [cell.strip() for cell in stripped.strip("|").split("|")]
            table_buffer.append(row)
            continue

        if table_buffer:
            flush_table(document, table_buffer)
            table_buffer = []

        if stripped.startswith("# "):
            continue

        add_markdown_line(document, line)

    if table_buffer:
        flush_table(document, table_buffer)

    return document


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
